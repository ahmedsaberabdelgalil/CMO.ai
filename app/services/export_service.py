"""
Document export helpers.

Turns agent output (marketing plans, brand reports) into downloadable Word
documents. Lightweight markdown handling: ``#``/``##`` headings, ``**bold**``,
and ``-``/``*`` bullets are mapped to native Word styles.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Render inline **bold** segments as bold runs."""
    pos = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx(*, title: str, content: str, subtitle: str | None = None) -> bytes:
    """Build a .docx document from a title and (lightly markdown) body text."""
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    document.add_heading(title or "CMO.AI Document", level=0)
    if subtitle:
        sub = document.add_paragraph()
        run = sub.add_run(subtitle)
        run.italic = True

    for raw_line in (content or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            document.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(heading_match.group(2).strip(), level=level)
            continue

        bullet_match = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_runs_with_bold(paragraph, bullet_match.group(1).strip())
            continue

        numbered_match = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if numbered_match:
            paragraph = document.add_paragraph(style="List Number")
            _add_runs_with_bold(paragraph, numbered_match.group(1).strip())
            continue

        paragraph = document.add_paragraph()
        _add_runs_with_bold(paragraph, line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
